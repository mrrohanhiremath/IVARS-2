from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_border(section):
    """Add a border to the page"""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '24')
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), '8B0000')
        pgBorders.append(border_el)
    
    sectPr.append(pgBorders)

def create_cocomo_document():
    doc = Document()
    
    # Set up the document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        add_page_border(section)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Incident Verification and Response System')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()  # Empty line
    
    # Main Heading
    heading = doc.add_heading('COCOMO Cost Estimation', level=1)
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.bold = True
    
    doc.add_paragraph()  # Empty line
    
    # Introduction paragraph
    intro = doc.add_paragraph(
        "Effort and cost estimation is a crucial activity in software project management, as it helps in "
        "planning resources, scheduling activities, and estimating the overall project budget. For the "
        "Incident Verification and Response System (IVARS), the Basic COCOMO (Constructive Cost Model) "
        "was used to estimate the development effort, time, team size, and cost."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in intro.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    para2 = doc.add_paragraph(
        "The estimation considers only the manually written source code, excluding external libraries, "
        "frameworks, and auto-generated files. Based on project characteristics such as moderate size, "
        "well-understood requirements, and a small experienced team, the Organic mode of the Basic "
        "COCOMO model was selected."
    )
    para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in para2.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    # Overview of Basic COCOMO Model
    doc.add_heading('Overview of Basic COCOMO Model', level=2)
    
    overview = doc.add_paragraph(
        "The Basic COCOMO model estimates software development effort and schedule primarily "
        "based on the size of the project measured in KLOC (Thousands of Lines of Code). It "
        "classifies projects into three types: Organic, Semi-detached, and Embedded."
    )
    overview.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in overview.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    para3 = doc.add_paragraph(
        "Since this project is a web-based application with moderate complexity and a flexible "
        "development environment, it falls under the Organic project category."
    )
    para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in para3.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    # Equations
    equations_heading = doc.add_paragraph()
    equations_heading.add_run('.The equations used are:').font.size = Pt(12)
    
    # Bullet points for equations
    equations = [
        "Effort (E) = a × (KLOC)ᵇ (Person-Months)",
        "Development Time (T) = c × (E)ᵈ (Months)",
        "Team Size = E / T",
        "Total Cost = Effort × Cost per Person-Month"
    ]
    
    for eq in equations:
        p = doc.add_paragraph(eq, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run.font.bold = True
    
    doc.add_paragraph()
    
    # Project Size Estimation
    doc.add_page_break()
    
    heading2 = doc.add_heading('Project Size Estimation', level=2)
    
    intro_text = doc.add_paragraph(
        "The total size of the project was calculated by summing the lines of code developed across "
        "different modules."
    )
    intro_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in intro_text.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Table 5.1: Project Size Estimation
    table1 = doc.add_table(rows=5, cols=3)
    table1.style = 'Table Grid'
    
    # Header row
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Module'
    hdr_cells[1].text = 'Technology Used'
    hdr_cells[2].text = 'Lines of Code (LOC)'
    
    # Make header bold and centered
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    # Data rows
    data = [
        ['Backend', 'Node.js / Express', '1,931'],
        ['Frontend', 'React / TypeScript', '2,807'],
        ['UI & Utilities', 'CSS / HTML', '1,006'],
        ['Total', '', '5,744 LOC']
    ]
    
    for i, row_data in enumerate(data, start=1):
        row_cells = table1.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
            row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in row_cells[j].paragraphs[0].runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                if i == 4:  # Total row
                    run.font.bold = True
    
    # Add another row for Project Size
    row = table1.add_row()
    row.cells[0].text = 'Project Size'
    row.cells[2].text = '5.7 KLOC'
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if cell != row.cells[0] else WD_ALIGN_PARAGRAPH.LEFT
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run.font.bold = True
    
    doc.add_paragraph()
    
    # Table caption
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run('Table 5.1: Project Size Estimation')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Page 3
    doc.add_page_break()
    
    # COCOMO Model Constants
    doc.add_heading('COCOMO Model Constants (Organic Mode)', level=2)
    
    constants_intro = doc.add_paragraph(
        "Based on the COCOMO reference values provided in the model documentation:"
    )
    for run in constants_intro.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Table 5.2: Constants
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Parameter'
    hdr_cells2[1].text = 'Value'
    
    for cell in hdr_cells2:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    # Constants data
    constants = [
        ['a', '2.4'],
        ['b', '1.05'],
        ['c', '2.5'],
        ['d', '0.38']
    ]
    
    for i, const_data in enumerate(constants, start=1):
        cells = table2.rows[i].cells
        cells[0].text = const_data[0]
        cells[1].text = const_data[1]
        for cell in cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    caption2 = doc.add_paragraph()
    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption2.add_run('Table 5.2: Organic Model Constants')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Effort Estimation
    doc.add_heading('Effort Estimation', level=2)
    
    effort_para = doc.add_paragraph(
        "Effort represents the total amount of work required to develop the software, measured in "
        "person-months."
    )
    effort_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in effort_para.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    calc_heading = doc.add_paragraph()
    run = calc_heading.add_run('Calculation:')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    calc1 = doc.add_paragraph()
    calc1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = calc1.add_run('E=2.4×(5.7)¹·⁰⁵')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    calc2 = doc.add_paragraph()
    calc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = calc2.add_run('E=14.89 Person-Months')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    result_heading = doc.add_paragraph()
    run = result_heading.add_run('Result:')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    result = doc.add_paragraph()
    run = result.add_run('Estimated Effort ≈ 15 Person-Months')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    # Development Time Estimation
    doc.add_heading('Development Time Estimation', level=2)
    
    time_para = doc.add_paragraph(
        "Development time indicates the total calendar time required to complete the project."
    )
    time_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in time_para.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    calc_heading2 = doc.add_paragraph()
    run = calc_heading2.add_run('Calculation:')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    calc3 = doc.add_paragraph()
    calc3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = calc3.add_run('T=2.5×(14.89)⁰·³⁸')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    calc4 = doc.add_paragraph()
    calc4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = calc4.add_run('T=6.79 months')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    time_result = doc.add_paragraph(
        "Considering parallel development and efficient task distribution, the effective development "
        "duration was approximately 5.3 months, which closely matches the actual project timeline."
    )
    time_result.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in time_result.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    # Page 4
    doc.add_page_break()
    
    # Team Size Estimation
    doc.add_heading('Team Size Estimation', level=2)
    
    team_intro = doc.add_paragraph(
        "The estimated number of developers required is calculated as:"
    )
    team_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in team_intro.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    team_calc = doc.add_paragraph()
    team_calc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = team_calc.add_run('Team Size=14.89/6.79 ≈2.19')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    result_heading2 = doc.add_paragraph()
    run = result_heading2.add_run('Result:')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    team_result = doc.add_paragraph()
    run = team_result.add_run('Estimated Team Size ≈ 3 members')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    actual_team = doc.add_paragraph(
        "The actual team consisted of 4 members, which ensured better workload sharing and timely "
        "completion."
    )
    actual_team.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in actual_team.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Cost Estimation
    doc.add_heading('Cost Estimation', level=2)
    
    cost_intro = doc.add_paragraph(
        "The total project cost is calculated based on the estimated effort and average cost per person-"
        "month."
    )
    cost_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in cost_intro.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    assumption = doc.add_paragraph()
    run = assumption.add_run('Assumption:')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    cost_assumption = doc.add_paragraph()
    run = cost_assumption.add_run('Average cost per person-month = ₹ 4.2 Lakhs')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    calc_heading3 = doc.add_paragraph()
    run = calc_heading3.add_run('Calculation:')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    cost_calc = doc.add_paragraph()
    cost_calc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cost_calc.add_run('Total Cost=15×4.2=63 Lakhs')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    result_heading3 = doc.add_paragraph()
    run = result_heading3.add_run('Result:')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    cost_result = doc.add_paragraph()
    run = cost_result.add_run('Estimated Project Cost ≈ ₹63 Lakhs')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Summary Table
    doc.add_heading('Summary of Estimation Results', level=2)
    
    # Table 5.3: Summary
    table3 = doc.add_table(rows=7, cols=2)
    table3.style = 'Table Grid'
    
    # Header
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'Parameter'
    hdr_cells3[1].text = 'Estimated Value'
    
    for cell in hdr_cells3:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    # Summary data
    summary_data = [
        ['Project Type', 'Organic'],
        ['Project Size', '5.7 KLOC'],
        ['Effort', '~15 Person-Months'],
        ['Development Time', '~5.3 Months'],
        ['Team Size', '~3 Members'],
        ['Estimated Cost', '~₹63 Lakhs']
    ]
    
    for i, sum_data in enumerate(summary_data, start=1):
        cells = table3.rows[i].cells
        cells[0].text = sum_data[0]
        cells[1].text = sum_data[1]
        for cell in cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if cell == cells[1] else WD_ALIGN_PARAGRAPH.LEFT
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    caption3 = doc.add_paragraph()
    caption3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption3.add_run('Table 5.3: COCOMO Estimation Summary')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Dept. of CSE,BITM,Ballari')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    page_num = doc.add_paragraph()
    page_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = page_num.add_run('Page 14')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Save the document
    doc.save('IVARS_COCOMO_Estimation.docx')
    print("✅ Document created successfully: IVARS_COCOMO_Estimation.docx")

if __name__ == "__main__":
    create_cocomo_document()
