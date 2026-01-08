from docx import Document

# Read the Word document
doc = Document(r"C:\Users\rohan\OneDrive\Documents\IVARS\IVARS-REPORT.docx")

print("=" * 80)
print("DOCUMENT STRUCTURE")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style = para.style.name
        text = para.text.strip()
        print(f"\n[Para {i}] Style: {style}")
        print(f"Text: {text[:100]}...")
        
print("\n" + "=" * 80)
print("Total paragraphs:", len(doc.paragraphs))
print("=" * 80)
